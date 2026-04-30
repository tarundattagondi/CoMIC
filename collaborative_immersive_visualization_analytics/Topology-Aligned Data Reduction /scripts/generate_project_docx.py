"""Generate a comprehensive DOCX report for the CIVA Reduction project.

Uses python-docx. Writes to: docs/CIVA_Reduction_Project_Report.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_ROOT / "docs" / "CIVA_Reduction_Project_Report.docx"


def set_cell_bg(cell, hex_color: str) -> None:
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    shading = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), hex_color)
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        for r in p.runs:
            r.font.size = Pt(11)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], header: tuple[str, str] = ("Field", "Value")) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    for k, v in rows:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    from docx.oxml.ns import qn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")


def build_doc() -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- Title page ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Topology-Aligned Data Reduction for Immersive Analytics of Structured Volumetric Data")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("CIVA Reduction — Final Project Report")
    r.italic = True
    r.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Sai Kiran Annam, Tarun Datta Gondi, Fahim Asrad Nafis\n"
        "Mobile and Immersive Computing, George Mason University\n"
        "April 2026"
    )

    doc.add_paragraph()

    # ---------- Abstract ----------
    add_heading(doc, "Abstract", level=1)
    add_para(
        doc,
        "Scientific workflows in simulation, imaging, and spatial analytics increasingly depend on high-resolution "
        "structured volumetric scalar data. While immersive analytics offers stronger depth perception and spatial "
        "understanding than flat-screen visualization, practical WebXR systems are constrained by strict frame-time "
        "budgets, browser/runtime overhead, and interaction-driven shifts in analytical focus. This project reframes "
        "reduction as an interaction-driven, topology-aligned, reversible, and measurable runtime behavior rather than "
        "a one-time preprocessing operation."
    )
    add_para(
        doc,
        "We implemented a complete browser-to-backend pipeline using React + TypeScript + VTK.js on the frontend and "
        "FastAPI + VTK/TTK on the backend. The system supports dynamic persistence-aware reduction, spatial LOD control, "
        "feature-focused contextual rendering (slice plane, contextual dimming, ROI wireframe, scalar rollback), explicit "
        "reduction-state modeling via a finite-state machine, and runtime observability through integrated metrics and "
        "session logging. A structured dataset descriptor pipeline and backend health/reduction endpoints provide "
        "robustness for repeatable system-study scenarios."
    )
    add_para(
        doc,
        "The platform demonstrates that topology-aligned reduction principles can be operationalized in WebXR-era "
        "systems without requiring novel reduction algorithms. The primary contribution is a working, modular, and "
        "reproducible system design that links interaction events to measurable reduction behavior."
    )

    add_para(doc, "Index Terms", bold=True)
    add_para(
        doc,
        "Immersive analytics, WebXR, topology-aligned reduction, volumetric visualization, VTK.js, TTK, level of detail, "
        "reproducible system study, interactive scientific visualization."
    )

    doc.add_page_break()

    # ---------- I. Introduction ----------
    add_heading(doc, "I. Introduction", level=1)
    add_para(
        doc,
        "Large structured volumetric datasets are central to scientific and engineering analysis. Domains such as "
        "computational fluid dynamics, climate modeling, and medical imaging routinely produce scalar fields on dense "
        "regular or rectilinear grids. Even one moderate-resolution volume can contain tens of millions of voxels, and "
        "modern workflows often involve multiple variables, time steps, and feature-specific exploration tasks. "
        "Rendering such data at full fidelity in real time is expensive even on desktop systems; in immersive contexts, "
        "the challenge is amplified by frame-rate and latency expectations required to preserve comfort."
    )
    add_para(
        doc,
        "Immersive analytics is attractive because it improves spatial understanding through embodied movement, "
        "stereoscopic depth cues, and direct spatial interaction. However, these strengths do not remove data "
        "complexity — they increase the need for dynamic abstraction, because a user’s viewpoint, task objective, and "
        "feature interest can change continuously. Static reduction pipelines are poorly aligned with this interaction-"
        "rich environment."
    )
    add_para(
        doc,
        "This project addresses that mismatch by treating reduction as a runtime system behavior tied to interaction "
        "state. We do not propose a new topological algorithm. Instead, we integrate topology-aligned principles into a "
        "practical web-native architecture and evaluate the resulting behavior through system-level instrumentation."
    )

    # ---------- II. Project Context ----------
    add_heading(doc, "II. Project Context and Prior Work", level=1)
    add_heading(doc, "A. Problem Context", level=2)
    add_numbered(
        doc,
        [
            "Performance bottlenecks: Full-resolution rendering is costly in browser pipelines and can destabilize frame timing.",
            "Visual clutter: Dense scalar fields can overwhelm interpretability when all detail remains visible.",
            "Static abstraction mismatch: Offline reduction cannot adapt to evolving interaction context during immersive analysis.",
        ],
    )

    add_heading(doc, "B. Topology-Aligned Reduction Motivation", level=2)
    add_para(
        doc,
        "Topological analysis provides a principled way to reason about feature significance in scalar fields. "
        "Persistence-guided simplification, contour-tree-informed abstractions, and structure-preserving filtering are "
        "preferable to naive uniform reduction when interpretability matters. Our project adopts this rationale at the "
        "system-design level: reduction choices should preserve meaningful structural behavior rather than merely "
        "decrease sample count."
    )

    add_heading(doc, "C. Baseline Before the Status Report", level=2)
    add_bullets(
        doc,
        [
            "Baseline VTK.js volume rendering",
            "Modular frontend folders (core, data, state, interaction, metrics)",
            "Initial LOD and persistence-threshold controls",
            "Backend scaffolding for dynamic reduction",
            "Early WebXR integration and feature detection",
            "Dataset descriptor foundations",
        ],
    )

    # ---------- III. Research Questions ----------
    add_heading(doc, "III. Final Objectives and Research Questions", level=1)
    add_bullets(
        doc,
        [
            "RQ1: How can topology-aligned reduction be integrated into immersive analytics as an interaction-driven system behavior, not static preprocessing?",
            "RQ2: What are the system-level performance and representational tradeoffs across interaction-driven reduction strategies in WebXR volumetric visualization?",
            "RQ3: How does immersive interaction support iterative, feature-centric reduction workflows for structured scalar fields?",
        ],
    )
    add_para(
        doc,
        "To answer these within project scope, we followed a system-study methodology emphasizing implementation "
        "correctness, runtime observability, reproducible traces, and engineering analysis rather than human-subject "
        "experimentation."
    )

    # ---------- IV. System Architecture ----------
    add_heading(doc, "IV. System Architecture", level=1)

    add_heading(doc, "A. Frontend Stack and Runtime Model", level=2)
    add_para(
        doc,
        "The frontend is implemented with React and TypeScript and uses VTK.js for volume rendering. The runtime "
        "model separates concerns into stable modules under src/:"
    )
    add_kv_table(
        doc,
        [
            ("src/core/renderer", "Render lifecycle, scene updates, camera loop, rendering validation."),
            ("src/core/webxr", "XR capability checks (xr-feature-detection), session manager, input manager."),
            ("src/data", "Descriptor loading, backend/static loading, async progress, LOD manager, feature manager, immutable assets."),
            ("src/state", "Centralized store, reduction FSM (reduction-fsm, reduction-phase), derived phase."),
            ("src/metrics", "FPS monitor, latency capture, memory hints, session event log, structured logging."),
            ("src/interaction", "Input/gesture handling, controller, ROI control primitive."),
            ("src/ui", "UI state and dashboard wiring (App.tsx owns the top-level view)."),
            ("src/config", "appConfig, config types, and validator for environment-driven settings."),
        ],
        header=("Module", "Responsibility"),
    )

    add_heading(doc, "B. Backend Reduction Service", level=2)
    add_para(doc, "The backend is a FastAPI service providing:")
    add_kv_table(
        doc,
        [
            ("GET /api/health", "{ ok, vtk, ttk, dataDir } — capability readiness."),
            ("GET /api/ttk", "Diagnostic environment detail about the TTK integration."),
            ("GET / POST /api/reduce", "Accepts datasetId, level (full|high|medium|low), optional persistenceThreshold. Returns reduced VTI bytes plus X-Reduction-Metadata header."),
        ],
        header=("Endpoint", "Purpose"),
    )
    add_para(
        doc,
        "The service sanitizes dataset IDs, resolves paths safely, and returns reduced VTI bytes with metadata in "
        "response headers (X-Reduction-Metadata) including reduction mode and output grid properties. This metadata is "
        "useful for traceability and post-hoc analysis."
    )

    add_heading(doc, "C. Data and Descriptor Layer", level=2)
    add_para(
        doc,
        "The shipped descriptor for ctBones captures dimensions, spacing, scalar fields, voxel count, and path. The "
        "default source volume is 256 × 256 × 256 with a single scalar field (ImageScalars) over range 0-255. Frontend "
        "loaders support progress callbacks and graceful fallback to static files when backend reduction is "
        "unavailable."
    )

    add_heading(doc, "D. Reduction-State Formalization (FSM)", level=2)
    add_para(doc, "Reduction phase is an explicit finite-state machine with canonical phases:")
    add_bullets(doc, ["idle", "base_volume", "lod_switched", "roi_refined", "feature_focus"])
    add_para(
        doc,
        "Phase is recomputed from current context rather than manually toggled, which improves determinism and "
        "reduces UI drift. A transition guard table supports strict event reasoning and clearer debugging. Key "
        "transitions include VOLUME_READY → base_volume, LOD_SELECT ≠ full → lod_switched, FEATURE_ENABLE → "
        "feature_focus, ROI_ENABLE → roi_refined, and symmetric disable transitions back to the appropriate prior "
        "state."
    )

    # ---------- V. Implementation Details ----------
    add_heading(doc, "V. Implementation Details and Final-Phase Progress", level=1)

    add_heading(doc, "A. Topology-Aligned Backend Pipeline Hardening", level=2)
    add_numbered(
        doc,
        [
            "Persistence-aware TTK integration: explicit persistence-diagram thresholding and topological simplification paths, with compatibility handling for wrapper/API variance.",
            "Scalar selection safety: input scalar selection logic avoids common TTK/VTK failures when arrays are not explicitly specified.",
            "Threshold mapping: normalized frontend persistence inputs can be mapped into dataset-specific absolute persistence scales using estimated maxima.",
            "Output normalization: reduction outputs are converted/validated into image data suitable for VTI writing.",
            "Structured logging: each reduction emits compact structured metadata for reproducible diagnostics.",
        ],
    )

    add_heading(doc, "B. LOD and Interaction-Coupled Reduction", level=2)
    add_para(doc, "LOD behavior now supports both manual and automatic modes:")
    add_bullets(
        doc,
        [
            "Manual selection among full, high, medium, low.",
            "Optional auto-LOD by camera distance using config-driven thresholds.",
            "ROI-triggered policy boost to higher LOD in selected situations.",
        ],
    )
    add_para(doc, "Given 256³ source resolution, this supports practical complexity scaling:")
    add_kv_table(
        doc,
        [
            ("Full / High", "256 × 256 × 256 = 16,777,216 voxels (baseline)."),
            ("Medium (2× shrink per axis)", "128 × 128 × 128 = 2,097,152 voxels (~8× reduction)."),
            ("Low (4× shrink per axis)", "64 × 64 × 64 = 262,144 voxels (~64× reduction)."),
        ],
        header=("LOD Level", "Voxel Count"),
    )

    add_heading(doc, "C. Feature-Focus Controls (Completed Additions)", level=2)
    add_numbered(
        doc,
        [
            "Slice plane clipping: a half-space clipping plane can be enabled through the volume center, with dimension (X/Y/Z) and offset control.",
            "Contextual dimming: the full volume can be dimmed to emphasize selected feature context without losing situational awareness.",
            "ROI wireframe preview: a world-space spherical ROI overlay can be toggled and resized for focus selection.",
            "Isosurface extraction and threshold-region rendering: feature operators driven by a configurable scalar threshold.",
            "Scalar field switching with rollback to previous active field.",
        ],
    )

    add_heading(doc, "D. Reversibility and Exploration Control", level=2)
    add_bullets(
        doc,
        [
            "Scalar rollback to previous field value.",
            "Global exploration reset restoring reduction, feature, ROI, and scalar state defaults.",
            "Session event append operations for every key action (timestamped, append-only).",
        ],
    )

    add_heading(doc, "E. WebXR Integration Quality", level=2)
    add_bullets(
        doc,
        [
            "Capability probing before session entry (xr-feature-detection).",
            "Session management through dedicated manager classes (xr-session-manager).",
            "Explicit input-session wiring (xr-input-manager).",
            "Clean exit and teardown behavior preserving state and instrumentation consistency across desktop and immersive modes.",
        ],
    )

    add_heading(doc, "F. Instrumentation and Reproducibility", level=2)
    add_numbered(
        doc,
        [
            "Performance monitor: rolling frame deltas with derived FPS mean/min/max and frame-time mean/p95.",
            "Latency capture: reduction/load latencies recorded by event kind (volume_load, LOD change, feature toggle, ROI refinement, threshold update).",
            "Memory hinting: JS heap snapshot approximation when the browser exposes performance.memory.",
            "Session event log: timestamped append-only event list.",
            "Exportable session package: downloadable JSON combining events and a metric snapshot.",
        ],
    )

    # ---------- Features Implemented table ----------
    add_heading(doc, "VI. Feature Inventory (What We Implemented)", level=1)
    add_para(doc, "The dashboard and state pipeline expose the following user-facing features:")
    add_kv_table(
        doc,
        [
            ("Dynamic LOD selection", "Manual switch across Full / High / Medium / Low; re-requests reduced VTI from backend."),
            ("Auto-LOD by camera distance", "Camera distance thresholds trigger LOD transitions automatically."),
            ("Persistence threshold slider", "Normalized [0,1] value sent to backend as persistenceThreshold; drives TTK simplification."),
            ("Slice plane clipping", "Enable/disable half-space clipping; choose dimension (X/Y/Z) and offset."),
            ("Contextual dimming", "Dims the whole volume to emphasize features of interest."),
            ("Isosurface / threshold-region rendering", "Feature operators that extract structures based on a scalar threshold."),
            ("ROI wireframe overlay", "Toggleable spherical ROI with adjustable radius and LOD-boost policy."),
            ("Scalar field switching", "Select the active scalar array; one-active model with rollback."),
            ("Scalar rollback", "Revert to the previously active scalar field."),
            ("Global exploration reset", "One-click reset of reduction, feature, ROI, and scalar state to defaults."),
            ("Reduction FSM", "Explicit phases (idle, base_volume, lod_switched, roi_refined, feature_focus) derived from context."),
            ("Performance dashboard", "Live FPS mean/min/max, frame-time mean/p95, sample count."),
            ("Reduction-latency capture", "Per-action latency recording (LOD, feature, ROI, threshold, volume load)."),
            ("Memory hint", "Approximate JS heap snapshot via performance.memory when available."),
            ("Session event log", "Timestamped append-only log of every key interaction."),
            ("Session export", "Downloadable JSON of events plus metric snapshot for reproducibility."),
            ("Backend health awareness", "UI reports TTK ready / VTK fallback / static fallback, and routes accordingly."),
            ("Reduction metadata header", "X-Reduction-Metadata returned per /api/reduce for traceability."),
            ("Dataset descriptor pipeline", "ctBones descriptor with dimensions, spacing, scalar fields, voxel count."),
            ("Async loader with progress callbacks", "Non-blocking VTI load with progress reporting and graceful fallback."),
            ("WebXR entry/exit with capability probing", "Safe immersive session lifecycle with desktop fallback."),
        ],
        header=("Feature", "Description"),
    )

    # ---------- VII. Tools and Technologies ----------
    add_heading(doc, "VII. Tools and Technologies Used", level=1)

    add_heading(doc, "A. Frontend", level=2)
    add_kv_table(
        doc,
        [
            ("Language", "TypeScript (strict mode, no implicit any, explicit return types)."),
            ("UI framework", "React 18."),
            ("Build tool", "Vite 5."),
            ("Volume rendering", "@kitware/vtk.js (v30.7.1)."),
            ("Immersive", "WebXR Device API (via VTK.js XR integration)."),
            ("Linting / format", "ESLint 9 (flat config), typescript-eslint, Prettier, eslint-config-prettier."),
            ("Module system", "ESM-only (no CommonJS)."),
            ("Runtime target", "Node.js 18.20.7 (pinned via .nvmrc and engines)."),
        ],
        header=("Layer", "Tool / Library"),
    )

    add_heading(doc, "B. Backend", level=2)
    add_kv_table(
        doc,
        [
            ("Language", "Python 3.11."),
            ("Web framework", "FastAPI."),
            ("ASGI server", "uvicorn[standard]."),
            ("Scientific IO / reduction", "VTK ≥ 9.3.0 (image I/O, spatial shrink)."),
            ("Topology analysis", "Topology ToolKit (TTK) via conda-forge (persistence-aware simplification)."),
            ("Fallback reduction path", "VTK-only spatial shrink when TTK is unavailable."),
            ("Packaging / runtime", "Dockerfile + docker-compose for Linux-container TTK (Apple Silicon friendly)."),
            ("Environment", "backend/.venv (venv) or conda env civa-backend."),
        ],
        header=("Layer", "Tool / Library"),
    )

    add_heading(doc, "C. Data and Formats", level=2)
    add_kv_table(
        doc,
        [
            ("Primary runtime format", ".vti (VTK Image Data) — structured scalar volume."),
            ("Primary dataset", "ctBones — 256 × 256 × 256 voxels, single scalar field ImageScalars, range 0-255."),
            ("Descriptor", "ctBones.descriptor.json — dimensions, spacing, scalar fields, voxel count, path."),
            ("Configuration", "src/config/appConfig.ts + .env (VITE_REDUCTION_API_URL, DATA_DIR)."),
        ],
        header=("Item", "Detail"),
    )

    add_heading(doc, "D. Developer Tooling", level=2)
    add_bullets(
        doc,
        [
            "Git / GitHub for version control (branch sai-tarun-code).",
            "VS Code workspace (.vscode/).",
            "TypeScript compiler (tsc) for type-check and build.",
            "Scripts: npm run dev / build / preview / lint / type-check / format.",
            "Backend helper scripts: backend/run.sh, backend/setup_ttk.sh, backend/run_ttk.sh.",
        ],
    )

    # ---------- VIII. End-to-End Flow ----------
    add_heading(doc, "VIII. End-to-End Flow and Design Diagrams (Described)", level=1)

    add_heading(doc, "A. End-to-End Architecture", level=2)
    add_para(
        doc,
        "User-in-browser/XR → React UI + Controls → State Store + FSM → VTK.js Renderer + Scene Manager → "
        "Metrics Collector + Session Log → Session JSON Export. State Store emits LOD/persistence requests to the "
        "FastAPI Reduction API, which reads the source dataset, applies TTK topological simplification and/or VTK LOD "
        "shrink, and returns reduced VTI bytes plus metadata to the renderer."
    )

    add_heading(doc, "B. Runtime Reduction as System Behavior", level=2)
    add_para(
        doc,
        "An interaction event (LOD, threshold, ROI, feature toggle) updates the FSM/store. If a data reload is needed, "
        "the system calls /api/reduce, TTK performs persistence simplification with the selected LOD, the reduced VTI "
        "is loaded into the scene, and the render loop runs. If no reload is needed (e.g., slice or threshold change), "
        "the feature pipeline is updated locally. In both cases, the render loop emits metrics and event-log entries."
    )

    add_heading(doc, "C. Observability: Action-to-Stable-Render Latency", level=2)
    add_para(
        doc,
        "On interaction, the UI calls beginActionLatency(kind); the scene applies the change (reload or local update); "
        "every frame records a delta; a stable-frame detector in the metrics collector determines when the scene has "
        "settled and records recordReductionLatency(kind_to_stable, ms). The result is a per-action latency classified "
        "by interaction kind, suitable for tradeoff analysis."
    )

    # ---------- IX. Evaluation ----------
    add_heading(doc, "IX. Evaluation Approach", level=1)
    add_heading(doc, "A. Methodological Framing", level=2)
    add_para(
        doc,
        "Consistent with proposal scope, evaluation is a system study rather than user study. The objective is to "
        "characterize how runtime reduction states influence performance and representational control under "
        "deterministic interactions."
    )

    add_heading(doc, "B. Scenario Structure", level=2)
    add_numbered(
        doc,
        [
            "Global context pass: begin at full/high context and observe baseline frame behavior.",
            "Distance-driven adaptation: enable auto-LOD and vary camera distance to trigger deterministic LOD transitions.",
            "Feature-focus pass: activate slicing and contextual dimming to isolate structural regions.",
            "ROI emphasis pass: enable ROI wireframe and adjust radius; observe coupled state and LOD behavior.",
            "Persistence sweep: vary persistence input, reload, and compare output metadata/performance traces.",
        ],
    )

    add_heading(doc, "C. Measurement Signals", level=2)
    add_bullets(
        doc,
        [
            "Frame-rate and frame-time distribution summaries (mean / min / max / p95).",
            "Load/reduction latency records per action kind.",
            "Voxel-scale and spacing metadata per loaded state.",
            "Action-sequenced event traces suitable for replay analysis.",
        ],
    )

    # ---------- X. Results ----------
    add_heading(doc, "X. Results", level=1)

    add_heading(doc, "A. Proposal-to-Implementation Completion", level=2)
    add_numbered(
        doc,
        [
            "Interaction-driven reduction — completed through live controls and backend coupling.",
            "Topology alignment — completed through persistence-driven TTK backend path.",
            "Reversibility — completed via rollback/reset behavior and explicit state derivation.",
            "Measurability — completed through integrated performance and session logging modules.",
            "Immersive deployment path — completed with WebXR feature detection and session lifecycle support.",
        ],
    )

    add_heading(doc, "B. Maturity Gains Since the Status Report", level=2)
    add_bullets(
        doc,
        [
            "Deterministic state machine formalization (reduction-fsm, reduction-phase).",
            "Completed metrics and export pipeline (performance-monitor, session-event-log, logging).",
            "Stronger backend diagnostics and metadata (X-Reduction-Metadata, /api/health, /api/ttk).",
            "Improved operational controls (auto-LOD, ROI coupling, feature-focus controls).",
        ],
    )

    add_heading(doc, "C. Tradeoff Analysis", level=2)
    add_numbered(
        doc,
        [
            "Performance vs detail: lower LOD reduces computational load but may suppress small structures.",
            "Persistence simplification vs fidelity: stronger thresholds remove low-significance topology but can hide subtle features.",
            "Feature emphasis vs context retention: clipping/dimming improves focus but reduces global context.",
            "Automation vs user control: auto-LOD improves responsiveness yet can reduce predictability if thresholds are poorly tuned.",
        ],
    )

    # ---------- XI. Discussion ----------
    add_heading(doc, "XI. Discussion: How the Final System Answers the RQs", level=1)
    add_heading(doc, "RQ1 (Interaction-Driven Integration)", level=2)
    add_para(
        doc,
        "Strongly addressed. Reduction is now a first-class runtime behavior connected to explicit state transitions, "
        "with direct ties to user actions and environment context. The architecture supports switching between "
        "reduction states without restarting the session or rebuilding offline assets."
    )
    add_heading(doc, "RQ2 (Performance-Representation Tradeoffs)", level=2)
    add_para(
        doc,
        "Addressed at the platform and instrumentation level. The system records enough data to analyze tradeoffs "
        "across LOD and persistence settings. While a larger matrix of benchmark runs would strengthen quantitative "
        "reporting, the current implementation provides a clear reproducible framework for those experiments."
    )
    add_heading(doc, "RQ3 (Iterative Feature-Centric Workflow)", level=2)
    add_para(
        doc,
        "Partially to strongly addressed. The iterative workflow is supported by slice/dim/ROI controls, rollback/"
        "reset behavior, and session event recording. ROI-localized high-resolution subvolume substitution and richer "
        "feature operators remain future extensions."
    )

    # ---------- XII. Lessons Learned ----------
    add_heading(doc, "XII. Lessons Learned", level=1)
    add_heading(doc, "A. Engineering", level=2)
    add_bullets(
        doc,
        [
            "State determinism is essential: deriving phase from context avoids subtle synchronization bugs.",
            "Health-aware integration matters: backend status endpoints are required for robust frontend behavior.",
            "Metrics should be native, not external: embedding instrumentation accelerates iteration and reproducibility.",
        ],
    )
    add_heading(doc, "B. Immersive Visualization", level=2)
    add_bullets(
        doc,
        [
            "Immersive rendering quality depends as much on stable frame behavior as on visual richness.",
            "Progressive reduction control is necessary for practical navigation and analysis in dense volumes.",
            "Feature-focused context controls are effective even before advanced geometry extraction is added.",
        ],
    )
    add_heading(doc, "C. Research Process", level=2)
    add_bullets(
        doc,
        [
            "In system studies, integration depth often determines value more than algorithm novelty.",
            "Reproducibility requires event logs and metadata from the beginning, not as a final add-on.",
            "Transparent UI reporting of backend/runtime state is important for trustworthy technical evaluation.",
        ],
    )

    # ---------- XIII. Limitations ----------
    add_heading(doc, "XIII. Limitations", level=1)
    add_numbered(
        doc,
        [
            "Dataset diversity: current validation is centered on a single canonical volume (ctBones).",
            "ROI refinement depth: ROI is currently a control/visualization primitive, not yet full localized data substitution.",
            "Feature operator breadth: limited threshold feature tooling relative to a production-grade extractor.",
            "Platform constraints: TTK deployment remains environment-dependent (notably on Apple Silicon).",
            "Benchmark breadth: expanded scenario runs across browsers/devices are still needed for a stronger empirical section.",
        ],
    )

    # ---------- XIV. Future Work ----------
    add_heading(doc, "XIV. Future Work", level=1)
    add_numbered(
        doc,
        [
            "Implement ROI-local high-resolution streaming/subvolume replacement with composited rendering.",
            "Add isosurface and threshold-region extraction as first-class feature operators with persistence-informed defaults.",
            "Expand latency metrics to classify action-to-stable-render per interaction type.",
            "Build scripted scenario runners for automated benchmark reproduction.",
            "Test multiple structured datasets and compare tuning robustness.",
            "Explore adaptive policies that jointly optimize FPS and feature-preservation heuristics.",
        ],
    )

    # ---------- XV. Build & Run ----------
    add_heading(doc, "XV. How to Build and Run", level=1)
    add_heading(doc, "A. Frontend", level=2)
    add_code_block(
        doc,
        "nvm use               # uses .nvmrc (Node 18.20.7)\n"
        "npm install\n"
        "npm run dev           # starts Vite dev server\n"
        "npm run build         # tsc -b && vite build\n"
        "npm run preview       # preview the built app\n"
        "npm run lint          # ESLint, zero warnings\n"
        "npm run type-check    # tsc --noEmit\n"
        "npm run format        # Prettier",
    )

    add_heading(doc, "B. Backend (venv, VTK fallback)", level=2)
    add_code_block(
        doc,
        "cd backend\n"
        "python -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install -r requirements.txt\n"
        "uvicorn main:app --reload --port 8000",
    )

    add_heading(doc, "C. Backend (conda, TTK enabled)", level=2)
    add_code_block(
        doc,
        "conda create -n civa-backend -c conda-forge topologytoolkit python=3.11 -y\n"
        "conda activate civa-backend\n"
        "pip install fastapi 'uvicorn[standard]'\n"
        "cd backend && uvicorn main:app --reload --port 8000\n"
        "# verify:\n"
        "curl -s http://localhost:8000/api/ttk\n"
        "curl -s http://localhost:8000/api/health",
    )

    add_heading(doc, "D. Backend (Docker, recommended on Apple Silicon)", level=2)
    add_code_block(
        doc,
        "# from project root\n"
        "docker compose up --build",
    )

    add_heading(doc, "E. Frontend → Backend wiring", level=2)
    add_code_block(doc, "# .env\nVITE_REDUCTION_API_URL=http://localhost:8000")

    # ---------- XVI. Project Structure ----------
    add_heading(doc, "XVI. Project Structure", level=1)
    add_code_block(
        doc,
        "CIVA_Reduction/\n"
        "├── src/\n"
        "│   ├── App.tsx                         # Top-level dashboard + controls\n"
        "│   ├── main.tsx                        # React entry\n"
        "│   ├── config/                         # appConfig, types, validator\n"
        "│   ├── core/\n"
        "│   │   ├── renderer/                   # renderer-manager, scene-manager, rendering-validation\n"
        "│   │   └── webxr/                      # xr-feature-detection, xr-session-manager, xr-input-manager\n"
        "│   ├── data/\n"
        "│   │   ├── data-loader.ts              # backend/static loader\n"
        "│   │   ├── async-loader.ts             # progress-aware loading\n"
        "│   │   ├── dataset-descriptor.ts       # ctBones descriptor integration\n"
        "│   │   ├── lod-manager.ts              # LOD orchestration\n"
        "│   │   ├── feature-manager.ts          # feature operators\n"
        "│   │   └── immutable-assets.ts\n"
        "│   ├── interaction/                    # controller, input/gesture, ROI primitive\n"
        "│   ├── metrics/\n"
        "│   │   ├── performance-monitor.ts      # FPS / frame-time / p95\n"
        "│   │   ├── metrics-collector.ts        # reduction/load latency\n"
        "│   │   ├── session-event-log.ts        # append-only log\n"
        "│   │   └── logging.ts                  # structured logging\n"
        "│   ├── state/\n"
        "│   │   ├── store.ts                    # central store\n"
        "│   │   ├── reduction-fsm.ts            # guard table + transitions\n"
        "│   │   └── reduction-phase.ts          # phase derivation from context\n"
        "│   ├── ui/                             # ui state\n"
        "│   └── types/                          # shared types\n"
        "├── backend/\n"
        "│   ├── main.py                         # FastAPI app, /api/health, /api/ttk, /api/reduce\n"
        "│   ├── reduce.py                       # TTK persistence + VTK shrink reduction pipeline\n"
        "│   ├── requirements.txt\n"
        "│   ├── Dockerfile, run.sh, run_ttk.sh, setup_ttk.sh\n"
        "├── data/datasets/                      # ctBones.vti + descriptor\n"
        "├── docs/                               # architecture, standards, final report\n"
        "├── scripts/                            # supporting scripts\n"
        "├── docker-compose.yml                  # Linux-container TTK deployment\n"
        "├── eslint.config.mjs, tsconfig.*.json, vite.config.ts, package.json\n"
        "└── README.md, LICENSE",
    )

    # ---------- XVII. Conclusion ----------
    add_heading(doc, "XVII. Conclusion", level=1)
    add_para(
        doc,
        "This final project demonstrates that topology-aligned data reduction can be operationalized as an "
        "interaction-driven and measurable behavior within a web-native immersive analytics system. By combining "
        "VTK.js/WebXR frontend design with a persistence-aware TTK backend and explicit reduction-state management, "
        "the implementation bridges an important gap between topology-informed visualization theory and practical "
        "immersive deployment."
    )
    add_para(
        doc,
        "The most significant final-phase achievements are not only feature additions but systems-level maturity "
        "gains: deterministic state modeling, robust backend observability, integrated quantitative instrumentation, "
        "and exportable session traces. These advances directly address the proposal’s core motivation and establish a "
        "defensible system-study artifact."
    )

    # ---------- References ----------
    add_heading(doc, "References", level=1)
    add_numbered(
        doc,
        [
            "J. Tierny, “Topology ToolKit: An Open-Source Library for Topological Data Analysis,” IEEE TVCG, vol. 26, no. 1, pp. 760-770, Jan. 2020.",
            "A. Gyulassy, P.-T. Bremer, B. Hamann, V. Pascucci, “A Practical Approach to Morse–Smale Complex Computation,” IEEE TVCG, vol. 14, no. 6, pp. 1619-1626, Nov.-Dec. 2008.",
            "M. Sedlmair, M. Meyer, M. Munzner, “Design Study Methodology: Reflections from the Trenches and the Stacks,” IEEE TVCG, vol. 18, no. 12, pp. 2431-2440, Dec. 2012.",
            "M. Satkowski, M. Sedlmair, M. Munzner, “Immersive Analytics: An Introduction,” IEEE CG&A, vol. 38, no. 2, pp. 16-18, Mar.-Apr. 2018.",
            "R. Cordeil, B. Bach, Y. Liu, et al., “Immersive Analytics,” IEEE CG&A, vol. 38, no. 2, pp. 66-79, Mar.-Apr. 2018.",
            "Topology ToolKit (TTK), Examples and Documentation. https://topology-tool-kit.github.io/examples/index.html",
            "CIVA Reduction project repository (src, backend, docs, data).",
        ],
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    main()
